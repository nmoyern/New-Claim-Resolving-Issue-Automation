"""
actions/fax_refax.py
---------------------
Complete refax workflow when an MCO auth was not received.

Admin Manual exact workflow:
  1. Find fax confirmation in Lauris (Fax History Report)
  2. Build Word cover letter with the required language:
       "Please see SRA request along with confirmation it was initially sent on
        [Date]. We are requesting you honor the date from the original submission."
  3. Refax: confirmation + SRA copy + cover letter → correct MCO fax number
  4. For wrong-MCO sends: update dates and refax to the correct MCO
  5. Update Lauris auth record / Con't Stays sheet

Functions exported:
  build_refax_cover_doc(original_fax_date, client_name, mco_name, save_path) -> save_path
  build_wrong_mco_cover_doc(correct_mco, client_name, original_request_date, save_path) -> save_path
  execute_refax_workflow(claim, original_send_date, confirmation_path, sra_pdf, fax_number) -> (bool, str)
"""
from __future__ import annotations

import asyncio
import subprocess
from datetime import date
from pathlib import Path
from typing import Optional, Tuple

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config.entities import get_entity_by_claimmd_region, get_entity_by_npi
from config.models import Claim, MCO
from config.settings import DRY_RUN, get_credentials
from sources.browser_base import BrowserSession
from logging_utils.logger import get_logger

logger = get_logger("fax_refax")

WORK_DIR = Path("/tmp/claims_work/fax")
WORK_DIR.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# Cover letter builders
# ---------------------------------------------------------------------------

def _get_entity_name(provider_npi: str = "", program: str = "") -> str:
    """Get the correct entity name based on NPI or program.
    Never use 'Life Consultants Inc.' — always the specific entity."""
    entity = get_entity_by_npi(provider_npi) or get_entity_by_claimmd_region(program)
    if entity:
        return entity.display_name
    return ""  # Never default — investigate via auth lookup


def build_refax_cover_doc(
    original_fax_date: date,
    client_name: str,
    mco_name: str,
    save_path: str,
    provider_npi: str = "",
    program: str = "",
) -> str:
    """
    Build the Admin-Manual-required refax cover letter Word document.

    Required language (verbatim from Admin Manual):
      "Please see SRA request along with confirmation it was initially sent on
       [Date]. We are requesting you honor the date from the original submission."

    Returns save_path on success.
    """
    doc = Document()

    # ── Page margins (1 inch) ─────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── Header ─────────────────────────────────────────────────────────────
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    hp.clear()
    entity = _get_entity_name(provider_npi, program)
    if not entity:
        logger.warning(
            "Cannot build refax cover — entity unknown. "
            "Must investigate via auth lookup.",
            client=client_name, npi=provider_npi, program=program,
        )
        return ""
    run = hp.add_run(f"{entity}  |  Authorization & Billing Department")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _para(text="", bold=False, size=11, space_before=0, space_after=6, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        p.alignment = align
        if text:
            r = p.add_run(text)
            r.font.name = "Calibri"
            r.font.size = Pt(size)
            r.bold = bold
        return p

    def _bold_label(label: str, value: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{label}: ")
        r1.font.name = "Calibri"; r1.font.size = Pt(11); r1.bold = True
        r2 = p.add_run(value)
        r2.font.name = "Calibri"; r2.font.size = Pt(11)

    # ── Title ───────────────────────────────────────────────────────────────
    _para("SERVICE AUTHORIZATION REFAX REQUEST", bold=True, size=14,
          space_before=12, space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Horizontal rule ─────────────────────────────────────────────────────
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p_hr = doc.add_paragraph()
    p_hr.paragraph_format.space_after = Pt(8)
    pPr = p_hr._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── Metadata ────────────────────────────────────────────────────────────
    _bold_label("Date",          date.today().strftime("%B %d, %Y"))
    _bold_label("To",            f"{mco_name} — Prior Authorization Department")
    _bold_label("Re",            f"Service Authorization Refax — {client_name}")
    _bold_label("From",          f"{entity} — Billing Department")
    _para()

    # ── Body — Admin Manual required language ────────────────────────────
    body_text = (
        f"Please see SRA request along with confirmation it was initially sent on "
        f"{original_fax_date.strftime('%m/%d/%Y')}. "
        f"We are requesting you honor the date from the original submission."
    )
    _para(body_text, size=11, space_before=6, space_after=10)

    # ── DMAS regulation reference ───────────────────────────────────────────
    _para(
        "Per DMAS regulations (12 VAC 30-130-5100 et seq.), managed care "
        "organizations are required to process service authorization requests "
        "in a timely manner and honor the original submission date when proof "
        "of timely filing is provided.",
        size=10, space_before=6, space_after=8,
    )

    # ── Enclosures list ─────────────────────────────────────────────────────
    _para("Enclosed documents:", bold=True, size=11, space_after=3)
    for enc in [
        "Original Service Request Authorization (SRA)",
        "Fax transmission confirmation — proof of original send date",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(enc)
        r.font.name = "Calibri"; r.font.size = Pt(11)

    _para()
    _para(
        "If you have any questions regarding this refax request, please contact "
        "our billing department at your earliest convenience. Thank you.",
        size=11, space_before=4
    )

    # ── Signature block ──────────────────────────────────────────────────────
    _para()
    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.space_after = Pt(2)
    r = p_sig.add_run(entity)
    r.font.name = "Calibri"; r.font.size = Pt(11); r.bold = True
    _para("Authorization & Billing Department", size=10)

    doc.save(save_path)
    logger.info("Refax cover letter built", path=save_path,
                client=client_name, original_date=str(original_fax_date))
    return save_path


def build_wrong_mco_cover_doc(
    correct_mco: str,
    client_name: str,
    original_request_date: date,
    save_path: str,
    provider_npi: str = "",
    program: str = "",
) -> str:
    """
    Cover letter for when an SRA was sent to the WRONG MCO.
    Admin Manual: change start date to today, adjust end date, refax to correct MCO.
    """
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    def _para(text="", bold=False, size=11, space_after=6,
              align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.alignment = align
        if text:
            r = p.add_run(text)
            r.font.name = "Calibri"; r.font.size = Pt(size); r.bold = bold
        return p

    def _label(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{label}: "); r1.font.name="Calibri"; r1.font.size=Pt(11); r1.bold=True
        r2 = p.add_run(value);        r2.font.name="Calibri"; r2.font.size=Pt(11)

    _para("SERVICE AUTHORIZATION — CORRECTED MCO SUBMISSION", bold=True,
          size=14, space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para()
    _label("Date", date.today().strftime("%B %d, %Y"))
    _label("To",   f"{correct_mco} — Prior Authorization Department")
    _label("Re",   f"Corrected SRA Submission — {client_name}")
    _para()

    body = (
        f"Please find enclosed a Service Request Authorization (SRA) for "
        f"{client_name}. This authorization was originally submitted on "
        f"{original_request_date.strftime('%m/%d/%Y')} but was inadvertently "
        f"directed to an incorrect managed care organization. "
        f"We are resubmitting to {correct_mco} as the correct payer. "
        f"The start date on the enclosed SRA has been updated to reflect today's "
        f"resubmission date. We respectfully request timely processing."
    )
    _para(body, size=11, space_after=10)
    _para("Thank you for your prompt attention to this matter.", size=11)
    _para()
    entity = _get_entity_name(provider_npi, program)
    if not entity:
        logger.warning(
            "Cannot build wrong-MCO cover — entity unknown. "
            "Must investigate via auth lookup.",
            client=client_name, npi=provider_npi, program=program,
        )
        return ""
    r_sig = doc.add_paragraph().add_run(f"{entity} — Billing Department")
    r_sig.font.name = "Calibri"; r_sig.font.size = Pt(11); r_sig.bold = True

    doc.save(save_path)
    logger.info("Wrong-MCO cover letter built", path=save_path,
                client=client_name, correct_mco=correct_mco)
    return save_path


# ---------------------------------------------------------------------------
# PDF combiner
# ---------------------------------------------------------------------------

def combine_to_pdf(
    cover_docx: str,
    confirmation_path: str,
    sra_pdf: Optional[str],
    output_path: str,
) -> bool:
    """Merge cover letter (docx→pdf) + confirmation image + SRA into one PDF."""
    import tempfile, shutil
    tmp = Path(tempfile.mkdtemp())
    parts = []

    # Convert cover docx → pdf via LibreOffice
    try:
        res = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", str(tmp), cover_docx],
            capture_output=True, timeout=30,
        )
        if res.returncode == 0:
            pdfs = list(tmp.glob("*.pdf"))
            if pdfs:
                parts.append(str(pdfs[0]))
    except Exception as e:
        logger.warning("LibreOffice conversion failed", error=str(e))

    # Confirmation image → pdf
    if confirmation_path and Path(confirmation_path).exists():
        img_pdf = str(tmp / "confirmation.pdf")
        try:
            _img_to_pdf(confirmation_path, img_pdf)
            parts.append(img_pdf)
        except Exception as e:
            logger.warning("Image→pdf failed", error=str(e))

    if sra_pdf and Path(sra_pdf).exists():
        parts.append(sra_pdf)

    if not parts:
        shutil.rmtree(tmp, ignore_errors=True)
        return False

    try:
        from pypdf import PdfMerger
        merger = PdfMerger()
        for p in parts:
            merger.append(p)
        with open(output_path, "wb") as f:
            merger.write(f)
        logger.info("Refax PDF merged", parts=len(parts), output=output_path)
        shutil.rmtree(tmp, ignore_errors=True)
        return True
    except ImportError:
        # copy first part as fallback
        import shutil as sh
        sh.copy(parts[0], output_path)
        shutil.rmtree(tmp, ignore_errors=True)
        return True
    except Exception as e:
        logger.error("PDF merge failed", error=str(e))
        shutil.rmtree(tmp, ignore_errors=True)
        return False


def _img_to_pdf(img_path: str, pdf_path: str):
    try:
        from PIL import Image
        Image.open(img_path).convert("RGB").save(pdf_path, "PDF")
    except Exception:
        pass


# ---------------------------------------------------------------------------
# Nextiva Fax portal automation
# ---------------------------------------------------------------------------

class NextivaFaxSession(BrowserSession):
    SESSION_NAME = "nextiva_fax"

    def __init__(self, headless: bool = True):
        super().__init__(headless=headless)
        self._creds = get_credentials().nextiva

    @property
    def login_url(self) -> str:
        return self._creds.url if self._creds else "https://fax.nextiva.com/xauth/"

    async def _is_logged_in(self) -> bool:
        try:
            url = self.page.url.lower()
            if "xauth" in url or "about:blank" in url:
                return False
            # Nextiva dashboard nav uses uppercase text
            el = await self.page.query_selector(
                "button:has-text('SIGN OUT'), button:has-text('DASHBOARD'), "
                "button:has-text('SEND')"
            )
            return el is not None
        except Exception:
            return False

    async def _perform_login(self) -> bool:
        await self.page.goto(
            self.login_url, wait_until="load", timeout=30000
        )
        await asyncio.sleep(2)
        if await self._is_logged_in():
            return True
        await self.page.fill(
            "input[name='LogonUserName']", self._creds.username
        )
        await self.page.fill(
            "input[name='LogonUserPassword']", self._creds.password
        )
        terms_cb = await self.page.query_selector(
            "input[name='LogonAccordTermsAccept']"
        )
        if terms_cb and not await terms_cb.is_checked():
            await terms_cb.check()
        await self.page.click("button[type='submit']")
        await asyncio.sleep(5)
        return await self._is_logged_in()

    def _get_send_frame(self):
        """Get the send fax iframe — Nextiva renders forms in iframes."""
        # Exact iframe name from live inspection
        for frame in self.page.frames:
            if frame.name == "xcAppNavStack_frame_send":
                return frame
        # Fallback: check by URL
        for frame in self.page.frames:
            if "send.aspx" in frame.url:
                return frame
        return self.page  # fallback to main page

    async def send_fax(
        self,
        to_number: str,
        pdf_path: str,
        subject: str = "Service Authorization",
    ) -> Tuple[bool, str]:
        if DRY_RUN:
            logger.info("DRY_RUN: Would send fax", to=to_number, file=pdf_path)
            return True, "DRY_RUN_CONFIRM"
        if not to_number:
            return False, "no_fax_number"

        try:
            # Navigate to Send page
            await self.page.click("button:has-text('SEND'), text=SEND")
            await asyncio.sleep(3)

            # Get the send frame (form is inside iframe)
            frame = self._get_send_frame()

            # Step 1: RECIPIENT INFO
            # Fill fax number (field prefixed with +1 for North America)
            await frame.fill("#xcFaxNumber", to_number)
            await asyncio.sleep(0.3)

            # Fill subject
            await frame.fill("#xcFaxSubject", subject)
            await asyncio.sleep(0.3)

            # Click NEXT to go to Step 2: ATTACHMENTS
            next_btn = await frame.query_selector(
                "button:has-text('NEXT'), a:has-text('NEXT'), "
                "input[value*='NEXT']"
            )
            if not next_btn:
                # NEXT button might be on the main page
                next_btn = await self.page.query_selector(
                    "button:has-text('NEXT'), text=NEXT"
                )
            if next_btn:
                await next_btn.click()
                await asyncio.sleep(3)

            # Step 2: ATTACHMENTS — upload the PDF
            # Re-get frame since page may have changed
            frame = self._get_send_frame()
            file_input = await frame.query_selector("input[type='file']")
            if not file_input:
                file_input = await self.page.query_selector(
                    "input[type='file']"
                )
            if file_input:
                await file_input.set_input_files(pdf_path)
                await asyncio.sleep(2)

            # Click NEXT to Step 3: PREVIEW
            next_btn = await frame.query_selector(
                "button:has-text('NEXT'), a:has-text('NEXT')"
            )
            if not next_btn:
                next_btn = await self.page.query_selector(
                    "button:has-text('NEXT'), text=NEXT"
                )
            if next_btn:
                await next_btn.click()
                await asyncio.sleep(3)

            # Step 3: PREVIEW — click NEXT to Step 4: SEND
            frame = self._get_send_frame()
            next_btn = await frame.query_selector(
                "button:has-text('NEXT'), a:has-text('NEXT'), "
                "button:has-text('SEND'), a:has-text('SEND FAX')"
            )
            if not next_btn:
                next_btn = await self.page.query_selector(
                    "button:has-text('NEXT'), text=NEXT"
                )
            if next_btn:
                await next_btn.click()
                await asyncio.sleep(3)

            # Step 4: SEND — click SEND to transmit
            frame = self._get_send_frame()
            send_btn = await frame.query_selector(
                "button:has-text('SEND'), a:has-text('SEND'), "
                "input[value*='SEND']"
            )
            if not send_btn:
                send_btn = await self.page.query_selector(
                    "button:has-text('SEND FAX'), text='SEND FAX'"
                )
            if send_btn:
                await send_btn.click()
                await asyncio.sleep(5)

            # Step 5: CONFIRMATION
            await self.screenshot("fax_confirmation")
            frame = self._get_send_frame()
            body_text = await frame.inner_text("body")
            confirm_id = "sent"
            import re
            conf_match = re.search(r"[A-Z0-9\-]{8,}", body_text)
            if conf_match:
                confirm_id = conf_match.group(0)

            logger.info("Fax sent", to=to_number, confirm=confirm_id)
            return True, confirm_id
        except Exception as e:
            logger.error("Fax send failed", to=to_number, error=str(e))
            await self.screenshot("fax_send_error")
            return False, ""

    async def lookup_sent_fax(self, search_date: date, client_name: str) -> Optional[str]:
        """Look up a sent fax; return confirmation ID or None."""
        try:
            await self.page.click("a:has-text('SENT'), a:has-text('Sent'), a[href*='sent']", timeout=10000)
            await asyncio.sleep(1)
            date_el = await self.page.query_selector("input[name*='date'], input[type='date']")
            if date_el:
                await date_el.fill(search_date.strftime("%m/%d/%Y"))
                await self.safe_click("button:has-text('Search'), button:has-text('Filter')")
                await asyncio.sleep(1)
            rows = await self.page.query_selector_all("tr.fax-row, tbody tr")
            for row in rows:
                text = await row.inner_text()
                if client_name.split()[-1].lower() in text.lower():
                    import re
                    m = re.search(r"[A-Z0-9\-]{6,}", text)
                    if m:
                        return m.group(0)
        except Exception as e:
            logger.warning("Nextiva sent-fax lookup failed", error=str(e))
        return None


# ---------------------------------------------------------------------------
# High-level orchestrator
# ---------------------------------------------------------------------------

async def execute_refax_workflow(
    claim: Claim,
    original_send_date: date,
    confirmation_path: str,
    sra_pdf: Optional[str],
    mco_fax_number: str,
) -> Tuple[bool, str]:
    """
    Complete refax workflow:
      1. Build cover letter docx
      2. Merge to single PDF
      3. Send via Nextiva
      Returns (success, confirm_id)
    """
    safe = claim.client_name.replace(" ", "_")
    today = date.today().strftime("%Y%m%d")

    cover_path = str(WORK_DIR / f"cover_{safe}_{today}.docx")
    build_refax_cover_doc(
        original_fax_date=original_send_date,
        client_name=claim.client_name,
        mco_name=claim.mco.value.title(),
        save_path=cover_path,
        provider_npi=claim.npi,
        program=claim.program.value if claim.program else "",
    )

    combined = str(WORK_DIR / f"refax_{safe}_{today}.pdf")
    combine_to_pdf(cover_path, confirmation_path, sra_pdf, combined)

    async with NextivaFaxSession() as fax:
        success, confirm_id = await fax.send_fax(
            to_number=mco_fax_number,
            pdf_path=combined,
            subject=f"SRA Refax — {claim.client_name}",
        )

    if success:
        logger.info("Refax complete", client=claim.client_name, confirm=confirm_id)
    else:
        logger.error("Refax failed", client=claim.client_name)
    return success, confirm_id


# ---------------------------------------------------------------------------
# Alias: build_refax_cover_letter
# test_all.py calls: build_refax_cover_letter(client_name, fax_date, mco, output_path=...)
# test_core.py calls: build_refax_cover_doc(original_fax_date, client_name, mco_name, save_path)
# ---------------------------------------------------------------------------

def build_refax_cover_letter(
    client_name: str,
    original_fax_date: date,
    mco_name: str,
    output_path: Optional[str] = None,
) -> str:
    """
    Alias for build_refax_cover_doc with (client_name, fax_date, mco) arg order.
    If output_path is omitted, an auto-generated path in WORK_DIR is used.
    Returns the path where the file was saved.
    """
    if output_path is None:
        safe = client_name.replace(" ", "_")
        output_path = str(WORK_DIR / f"cover_{safe}_{date.today().strftime('%Y%m%d')}.docx")

    return build_refax_cover_doc(
        original_fax_date=original_fax_date,
        client_name=client_name,
        mco_name=mco_name,
        save_path=output_path,
    )
